"""
IOCL PDF-to-Word Converter Project
-----------------------------------
DOCUMENT GENERATOR

This is the core piece of your project. It does two things:
  1. Takes a PDF and runs it through Docling to extract structured content
     (headings, paragraphs, tables) as Markdown text.
  2. Reads that Markdown line by line and builds a real, editable Word
     (.docx) file using python-docx.

This file has two main functions:
  - extract_with_docling(pdf_path)      -> gets structured text from the PDF
  - markdown_to_word(markdown_text, ..) -> turns that text into a .docx file

And one "master" function that does both steps together:
  - convert_pdf_to_word(pdf_path, output_path)

Install requirements (if not already done):
    pip install docling python-docx
"""

from docx import Document
from docx.shared import Pt


def extract_with_docling(pdf_path):
    """
    STEP 1: Read the PDF using Docling.

    Docling looks at the PDF, understands its layout (headings, paragraphs,
    tables), and gives us back that structure as Markdown text - the same
    kind of output you already saw in your comparison test.
    """
    from docling.document_converter import DocumentConverter

    converter = DocumentConverter()
    result = converter.convert(pdf_path)
    markdown_text = result.document.export_to_markdown()
    return markdown_text


def markdown_to_word(markdown_text, output_path):
    """
    STEP 2: Turn Docling's Markdown output into an actual Word file.

    We go through the text one line at a time and decide what each line is:
      - A line starting with "## "  -> it's a heading
      - A line starting with "|"    -> it's part of a table
      - Anything else (non-empty)   -> it's a normal paragraph

    Then we tell python-docx to add the matching element to the Word file.
    """
    document = Document()

    # Split the markdown into individual lines so we can look at them one by one
    lines = markdown_text.split("\n")

    # This will temporarily hold rows of a table while we're inside one
    table_rows_buffer = []

    def flush_table_buffer():
        """
        Helper function: if we were collecting table rows, build the actual
        Word table now and empty the buffer. This runs whenever we finish
        reading a table (i.e. the next line is NOT a table row anymore).
        """
        if not table_rows_buffer:
            return  # nothing to do, we weren't inside a table

        num_columns = len(table_rows_buffer[0])
        word_table = document.add_table(rows=0, cols=num_columns)
        word_table.style = "Light Grid Accent 1"  # a clean built-in Word table style

        for row_index, row_data in enumerate(table_rows_buffer):
            row_cells = word_table.add_row().cells
            for col_index, cell_text in enumerate(row_data):
                row_cells[col_index].text = cell_text.strip()
                # Make the header row (first row) bold
                if row_index == 0:
                    for paragraph in row_cells[col_index].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True

        table_rows_buffer.clear()
        document.add_paragraph("")  # small spacing gap after the table

    for line in lines:
        stripped_line = line.strip()

        # Skip the markdown "separator" row that looks like |---|---|---|
        if stripped_line.startswith("|") and set(stripped_line.replace("|", "").replace("-", "").strip()) == set():
            continue

        # Case 1: this line is part of a table
        if stripped_line.startswith("|"):
            # Split "| A | B | C |" into ["A", "B", "C"]
            cells = [cell.strip() for cell in stripped_line.strip("|").split("|")]
            table_rows_buffer.append(cells)
            continue
        else:
            # We've reached a non-table line, so finish any table we were building
            flush_table_buffer()

        # Case 2: this line is a heading (Docling uses "## " for headings)
        if stripped_line.startswith("## "):
            heading_text = stripped_line.replace("## ", "")
            document.add_heading(heading_text, level=1)

        # Case 3: this line is a top-level title (Docling sometimes uses "# ")
        elif stripped_line.startswith("# "):
            heading_text = stripped_line.replace("# ", "")
            document.add_heading(heading_text, level=0)

        # Case 4: empty line - just skip it, don't add blank paragraphs everywhere
        elif stripped_line == "":
            continue

        # Case 5: anything else is a normal paragraph
        else:
            paragraph = document.add_paragraph(stripped_line)
            paragraph.style.font.size = Pt(11)

    # In case the document ends while we were still inside a table
    flush_table_buffer()

    document.save(output_path)
    print(f"Word file created: {output_path}")


def convert_pdf_to_word(pdf_path, output_path):
    """
    MASTER FUNCTION: runs both steps together.
    This is the one function your FastAPI backend will eventually call.
    """
    print(f"Reading and extracting structure from: {pdf_path}")
    markdown_text = extract_with_docling(pdf_path)

    print("Building the Word document...")
    markdown_to_word(markdown_text, output_path)

    print("Done!")


if __name__ == "__main__":
    # Change these two lines to match your actual file names
    input_pdf = "sample_multipage_bill.pdf"
    output_docx = "converted_multipage_bill.docx"

    convert_pdf_to_word(input_pdf, output_docx)